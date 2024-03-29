import os

from OpenAIGPT import OpenAIGPT
from NLTK_tokenizer import NLTK_Tokenizer

import pandas as pd

from tika import parser  # for reading pdf
import docx
import openai
import time
import logging
import json
import re
# Use your own API key
from api_import import api_import
api_import()

import io
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage


from tkinter import *
from tkinter import ttk
import requests, math
from requests.structures import CaseInsensitiveDict


# use this to read unreadable pdfs
from PIL import Image
import pytesseract
import pdf2image


class Text_Miner():
    def __init__(self, root, mode):

        # apikey_location = r"C:\Users\d.los\OneDrive - Berenschot\Bureaublad\chatgpt openai key.txt"
        # apikey_location = r"C:\Users\danie\OneDrive\Bureaublad\Coding\api keys\openai key.txt"
        # with open(apikey_location) as f:
        #     self.key = f.readline()
        # openai.api_key = self.key

        # self.root = r"C:\Users\d.los\OneDrive - Berenschot\Documenten\testdocs"
        # self.root = r"C:\Users\d.los\OneDrive - Berenschot\Bureaublad\test ai"
        # self.root = "test documenten"
        self.root = root
        self.name = input('Verzin een naam voor de data: ')

        # the base of the folder that you want to siff through

        self.supported_languages = pytesseract.get_languages()
        self.target_language = 'nld'
        self.langs = {'nld': 'dutch', 'eng': 'english'}
        self.prompt = str() # the actual prompt that will be sent to the ai
        # self.mode = str('Noem alle maatregelen uit de volgende tekst die te maken hebben met verbeteren van de luchtkwaliteit op commagescheiden manier' +
        #             # 'en een schatting per maatregel of deze in de fase: "voornemen", "beginfase", "uitvoering" of "geimplementeerd" is' +
        #                 " als er geen maatregelen zijn genoemd, antwoord dan - . " +
        #                 " scheid de volgende twee aanvullingen met een |, ze moeten later naar een kolom worden verwerkt "
        #                 " Maak per bullet zelf een inschatting met welk overkoepelend thema het te maken heeft en "
        #                 "als er iets vermeld staat over voortgang meld dat dan ook " +
        #
        #                 ""
        #                 )

        self.mode = mode
        # self.name = str()

        # TODO: specify modes that this thing can operate with

        # the list of documents for every file in the root
        self.doclist = {}
        self.doclist_short = {}

        # list of all strings in all docments per dict.
        self.stringdict = {}
        self.outputdict = {}

        self.df = pd.DataFrame()

        # make an unique filename
        # self.file_name = str("Output " + str(time.strftime("%m %d %H%M%S ")) + ".json")

        self.process_speed = 50  # docs per minute
        self.max_query_length = 1000

        self.output = {}
        self.estimated_tokencount = False
        self.estim_costs = {'GPT3.5 Turbo' : False, }#'davinci': False, 'ada': False, 'GPT4-8k' : False} # values are false for now, updated later
        self.accord = False
        self.AI = OpenAIGPT()

    def get_languages(self, **kwargs):
        # this function checks if the target language is possible to work with.
        print("languages that we can work with: ", (self.supported_languages))
        try:
            if self.target_language in self.supported_languages:
                print("The target language ", self.target_language, " is in the supported languages")
        except:
            print('please use a string and check the shorthand version')

    def get_structure(self):
        # this fucntion reads in all the filenames, as well as order them per extension into a dict.
        self.doclist = {}
        self.doclist_short = {}
        for root, place, documents in os.walk(self.root):

            for document in documents:
                print(document)
                # add extension name into dict
                name, extension = document.split(".", 1)
                if extension not in self.doclist.keys():
                    self.doclist[extension] = []
                    self.doclist_short[extension] = []

                # add document file location to the dict
                self.doclist[extension].append(root + '\\' + document)
                # add document name to the sort dict
                self.doclist_short[extension].append(document)
                # create a dictionary per document name so the text can be put in there
                # self.stringlist[name] = []



    def add_ocr(self):
        # TODO: add a thing that adds ocr with pytesseract to the pdfs so images and unselectable pdf can be loaded in
        pass
        # self.stringlist[pdfpath].append(pytesseract.image_to_string(
        #       pdfpath,
        #       lang=self.target_language,
        #       config='--psm 6')
        #       )                                                                 # TODO: this function is still broken

    def read_files(self):
        # this is supposed to read all files that have a sensible extension (currently: doc, (selectable) pdf)


        def gettext(filename):
            # gets text from a docx file
            doc = docx.Document(filename)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            return '\n'.join(fullText)

        def convert_pdf_to_txt(path):
            resource_manager = PDFResourceManager()
            out_file = io.StringIO()
            # codec = 'utf-8'
            laparams = LAParams()
            device = TextConverter(resource_manager, out_file, laparams=laparams)
            fp = open(path, 'rb')
            interpreter = PDFPageInterpreter(resource_manager, device)
            for page in PDFPage.get_pages(fp, check_extractable=True):
                interpreter.process_page(page)
            fp.close()
            device.close()
            text = out_file.getvalue()
            out_file.close()
            return text

        def scrubstring(text):
            text = bytes(text, 'utf-8').decode('utf-8', "replace")
            text = text.replace("\n", " ")
            text = text.replace("\t", " ")
            text = ' '.join(text.split())

            return text

        doccount = 0
        charcount = 0
        for extension in self.doclist.keys():

            if extension == 'pdf':
                for item in self.doclist[extension]:
                    name = item.split('\\', -1)[-1]
                    text = convert_pdf_to_txt(item)
                    text = scrubstring(text)

                    self.stringdict.setdefault(name, [])
                    self.stringdict[name].append(text)

            if extension == 'docx':
                for item in self.doclist[extension]:
                    # counts how many files are processed
                    doccount += 1
                    print(doccount)
                    # gets the name of the file
                    name = item.split('\\', -1)[-1]
                    # extracts data
                    text = gettext(item)
                    # stores the string in the stringdict
                    text = scrubstring(text)

                    self.stringdict.setdefault(name, [])
                    self.stringdict[name].append(text)

            if extension == 'xlsx':
                print('xlsx files are not (yet) supported')

            else:
                print('There are files with un-integrated extensions, please check the folder.')
                doccount += 1
            # print(self.stringdict)

    def estimate_costs(self):
        '''
        This function creates an estimate of the costs by examining the text loaded in.
        It counts the amount of text by breaking it up into chuncks of length 1000 tokens (self.max_query_length)
        Then it multiplies this length with the costs per 1000 tokens as of 02-2023
        TODO: limitation is that gpt4 has different costs for the 32.000 token model and 8.000 token model
        '''
        length = 0
        for key, value in self.stringdict.items():
            length += len(value)

        queries = length // 4000

        seconds = queries * self.process_speed / 60
        seconds = round(seconds, 2)
        print(f'Iterations needed: {length}')
        # print(f'Estimated time needed for free version is {seconds}')

        tokencount = 0
        doccount = 0
        for string in self.stringdict.values():
            # print(string[0])
            prompt_to_inspect = string[0]
            doccount +=1
            print(doccount)
            tokencount += len(NLTK_Tokenizer(prompt_to_inspect, 1000))

        # update the user on costs
        print(f"\nThe amount of chuncks are {tokencount}")

        # self.estim_costs['GPT4-8k'] = str(round(tokencount * 0.06, 2)) + ' EUR'

        # self.estim_costs['davinci'] = str(round(tokencount * 0.02,2)) +  ' EUR'

        # self.estim_costs['ada'] = str(round(tokencount * 0.0004, 2)) + ' EUR'

        self.estim_costs['GPT3.5 Turbo'] = str(round(tokencount * 0.0002, 2)) + ' EUR'

        self.estimated_tokencount = tokencount

        return self.estim_costs

    # Functions and Objects
    def agree(self):
        if self.estimated_tokencount:
            print(f"do you agree with the estimated costs of: {self.estim_costs}")
            userinput = input("yes or no")
            if userinput.lower() == "yes":
                print("Input was yes, continuing the process")
                self.accord = True
            else:
                print("Input was not yes \n Terminating process...")
        else:
            print("No cost estimate has been done, please run self.estimate_costs()")


    def AI_interact(self):
        '''
        This section actually loads the text into openai
        :return:
        '''
        for docname, text in self.stringdict.items():
            # generate_text_with_prompt splits the prompt into multiple sections if too long
            # then it gets new data from the chatGPT

            output = self.AI.generate_text_with_prompt(prompt=text, mode=str(self.mode))

            self.outputdict.setdefault(docname, [])
            self.outputdict[docname].append(output)


    def write_to_doc(self): # TODO: for some reason this is not consistent
        ''' This writes the queries that were done by openai to a document '''

        if 'bulletpoints' in self.mode or 'bullet points' in self.mode:
            # This loop splits the sections of the text for bulletpoints
            rows = []
            for document_name, string in self.outputdict.items():
                elements = string[0].split('\n- ')
                for element in elements:
                    row = [document_name, element]
                    rows.append(row)


        doc = docx.Document()

        for document_name, string in self.outputdict.items():

            doc.add_paragraph(document_name)
            doc.add_paragraph(string)

        doc.save('output/' + self.name + 'doc100.docx')


    def write_to_xl(self, string):
        # Create a dataframe from the list of lists with column names

        bullets = self.AI.to_bullets(string)
        rows = bullets.split('\n- ')

        self.df = pd.DataFrame(rows, columns=['Document Name', 'Element'])

        self.df.to_excel('output/' + self.name + '.xlsx')
        # View the dataframe
        print(self.df)

    def open_file(self):
        ''' Open Json files '''
        with open(self.file_name) as json_file:
            data = json.load(json_file)

